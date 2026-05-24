from __future__ import annotations


import pytest

from compact_rag.common.exceptions import DocumentLoadError, UnsupportedFormatError
from compact_rag.ingestion.loader import (
    HTMLLoader,
    LoadedPage,
    LoaderFactory,
    MarkdownLoader,
    PDFLoader,
    TextLoader,
    _clean_pdf_text,
    BaseLoader,
    DOCXLoader,
)


# ── LoaderFactory ───────────────────────────────────────────────


class TestLoaderFactory:
    def test_returns_pdf_loader_for_pdf_extension(self):
        loader = LoaderFactory.get_loader("test.pdf")
        assert isinstance(loader, PDFLoader)

    def test_returns_pdf_loader_for_PDF_uppercase(self):
        loader = LoaderFactory.get_loader("TEST.PDF")
        assert isinstance(loader, PDFLoader)

    def test_returns_docx_loader(self):
        loader = LoaderFactory.get_loader("document.docx")
        assert isinstance(loader, DOCXLoader)

    def test_returns_text_loader(self):
        loader = LoaderFactory.get_loader("notes.txt")
        assert isinstance(loader, TextLoader)

    def test_returns_markdown_loader(self):
        loader = LoaderFactory.get_loader("README.md")
        assert isinstance(loader, MarkdownLoader)

    def test_returns_html_loader_for_html(self):
        loader = LoaderFactory.get_loader("index.html")
        assert isinstance(loader, HTMLLoader)

    def test_returns_html_loader_for_htm(self):
        loader = LoaderFactory.get_loader("index.htm")
        assert isinstance(loader, HTMLLoader)

    def test_raises_unsupported_format_for_unknown_extension(self):
        with pytest.raises(UnsupportedFormatError):
            LoaderFactory.get_loader("data.csv")

    def test_raises_unsupported_format_for_no_extension(self):
        with pytest.raises(UnsupportedFormatError):
            LoaderFactory.get_loader("Makefile")


# ── BaseLoader._get_file_info ────────────────────────────────────


class TestBaseLoaderGetFileInfo:
    def test_returns_correct_file_size(self, tmp_path):
        f = tmp_path / "hello.txt"
        f.write_text("hello world")
        info = BaseLoader._get_file_info(str(f))
        assert info["filename"] == "hello.txt"
        assert info["file_type"] == "txt"
        assert info["file_size"] == len("hello world")
        assert "hash" in info
        assert len(info["hash"]) == 64

    def test_hash_is_stable_for_same_content(self, tmp_path):
        f = tmp_path / "dup.txt"
        f.write_text("same content")
        info1 = BaseLoader._get_file_info(str(f))
        info2 = BaseLoader._get_file_info(str(f))
        assert info1["hash"] == info2["hash"]

    def test_hash_differs_for_different_content(self, tmp_path):
        f1 = tmp_path / "a.txt"
        f2 = tmp_path / "b.txt"
        f1.write_text("aaa")
        f2.write_text("bbb")
        assert BaseLoader._get_file_info(str(f1))["hash"] != BaseLoader._get_file_info(str(f2))["hash"]

    def test_unknown_extension_uses_raw_suffix(self, tmp_path):
        f = tmp_path / "data.xyz"
        f.write_text("content")
        info = BaseLoader._get_file_info(str(f))
        # lstrip('.') strips the leading dot
        assert info["file_type"] == "xyz"


# ── PDFLoader ────────────────────────────────────────────────────


class TestPDFLoader:
    @pytest.mark.asyncio
    async def test_loads_pdf_and_returns_pages(self, tmp_path, mocker):
        """Test that a valid PDF returns list[LoadedPage] with extracted text."""
        from pypdf import PdfWriter

        pdf_path = tmp_path / "test.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        writer.add_blank_page(width=612, height=792)
        writer.write(pdf_path)

        # Minimal text stub — pypdf empty pages return ""
        # We mock _extract_with_pypdf to inject content
        loader = PDFLoader()
        mocker.patch.object(
            loader,
            "_extract_with_pdfplumber",
            return_value=None,
        )
        mocker.patch.object(
            loader,
            "_extract_with_pypdf",
            return_value=["Page one content", "Page two content"],
        )

        pages = await loader.load(str(pdf_path))
        assert isinstance(pages, list)
        assert len(pages) == 2
        assert all(isinstance(p, LoadedPage) for p in pages)
        assert pages[0].page_number == 1
        assert pages[1].page_number == 2
        assert "Page one content" in pages[0].content
        assert "Page two content" in pages[1].content
        assert pages[0].metadata.get("file_type") == "pdf"

    @pytest.mark.asyncio
    async def test_handles_corrupt_pdf(self, tmp_path):
        pdf_path = tmp_path / "corrupt.pdf"
        pdf_path.write_bytes(b"not a valid pdf file")

        loader = PDFLoader()
        with pytest.raises(DocumentLoadError):
            await loader.load(str(pdf_path))

    @pytest.mark.asyncio
    async def test_handles_missing_file(self, tmp_path):
        # _get_file_info is called before the try/except, so built-in FileNotFoundError
        loader = PDFLoader()
        with pytest.raises(FileNotFoundError):
            await loader.load(str(tmp_path / "nonexistent.pdf"))

    @pytest.mark.asyncio
    async def test_strips_cleaned_text(self, tmp_path, mocker):
        """Verify that _clean_pdf_text is applied to extracted content."""
        pdf_path = tmp_path / "test.pdf"
        pdf_path.write_bytes(b"%PDF-1.4")

        loader = PDFLoader()
        mocker.patch.object(loader, "_extract_with_pdfplumber", return_value=None)
        mocker.patch.object(loader, "_extract_with_pypdf", return_value=["(cid:123)  hello   world  "])

        pages = await loader.load(str(pdf_path))
        assert "(cid:123)" not in pages[0].content
        assert "hello world" in pages[0].content


# ── DOCXLoader ───────────────────────────────────────────────────


@pytest.mark.slow
class TestDOCXLoader:
    @pytest.mark.asyncio
    async def test_loads_docx_and_returns_pages(self, tmp_path):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("First paragraph with some content.")
        doc.add_paragraph("Second paragraph of the document.")
        doc.save(str(docx_path))

        loader = DOCXLoader()
        pages = await loader.load(str(docx_path))

        assert isinstance(pages, list)
        assert len(pages) >= 1
        assert all(isinstance(p, LoadedPage) for p in pages)
        assert "First paragraph" in pages[0].content
        assert "Second paragraph" in pages[0].content
        assert pages[0].metadata.get("file_type") == "docx"

    @pytest.mark.asyncio
    async def test_docx_with_table(self, tmp_path):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        docx_path = tmp_path / "tabled.docx"
        doc = Document()
        doc.add_paragraph("Before table.")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"
        doc.add_paragraph("After table.")
        doc.save(str(docx_path))

        loader = DOCXLoader()
        pages = await loader.load(str(docx_path))
        assert pages[0].tables
        assert "A" in pages[0].tables[0]
        assert pages[0].metadata.get("table_count") == 1

    @pytest.mark.asyncio
    async def test_handles_missing_file(self, tmp_path):
        # _get_file_info is called before the inner try/except
        loader = DOCXLoader()
        with pytest.raises(FileNotFoundError):
            await loader.load(str(tmp_path / "nonexistent.docx"))


# ── TextLoader ──────────────────────────────────────────────────


class TestTextLoader:
    @pytest.mark.asyncio
    async def test_reads_text_file_and_splits_into_pages(self, tmp_path):
        path = tmp_path / "sample.txt"
        content = "Line one.\nLine two.\nLine three.\n"
        path.write_text(content, encoding="utf-8")

        loader = TextLoader()
        pages = await loader.load(str(path))

        assert len(pages) >= 1
        assert "Line one" in pages[0].content
        assert pages[0].metadata.get("file_type") == "txt"

    @pytest.mark.asyncio
    async def test_handles_empty_file(self, tmp_path):
        path = tmp_path / "empty.txt"
        path.write_text("", encoding="utf-8")

        loader = TextLoader()
        pages = await loader.load(str(path))
        assert len(pages) == 1

    @pytest.mark.asyncio
    async def test_handles_unicode_content(self, tmp_path):
        path = tmp_path / "unicode.txt"
        path.write_text("你好世界 🌍 日本語テスト", encoding="utf-8")

        loader = TextLoader()
        pages = await loader.load(str(path))
        assert "你好世界" in pages[0].content
        assert "🌍" in pages[0].content

    @pytest.mark.asyncio
    async def test_falls_back_to_latin1_for_non_utf8(self, tmp_path):
        path = tmp_path / "latin1.txt"
        path.write_bytes("café ñoño naïve".encode("latin-1"))

        loader = TextLoader()
        pages = await loader.load(str(path))
        assert "café" in pages[0].content

    @pytest.mark.asyncio
    async def test_handles_missing_file(self, tmp_path):
        # _get_file_info is called before the try/except
        loader = TextLoader()
        with pytest.raises(FileNotFoundError):
            await loader.load(str(tmp_path / "nonexistent.txt"))

    @pytest.mark.asyncio
    async def test_large_text_splits_into_multiple_pages(self, tmp_path):
        path = tmp_path / "large.txt"
        content = "x" * 6000
        path.write_text(content, encoding="utf-8")

        loader = TextLoader()
        pages = await loader.load(str(path))
        assert len(pages) >= 2
        assert pages[0].metadata["page_count"] == len(pages)


# ── MarkdownLoader ──────────────────────────────────────────────


class TestMarkdownLoader:
    @pytest.mark.asyncio
    async def test_reads_markdown(self, tmp_path):
        path = tmp_path / "doc.md"
        path.write_text("# Title\n\nParagraph text here.", encoding="utf-8")

        loader = MarkdownLoader()
        pages = await loader.load(str(path))

        assert len(pages) >= 1
        assert "# Title" in pages[0].content
        assert pages[0].metadata.get("file_type") == "md"

    @pytest.mark.asyncio
    async def test_handles_empty_file(self, tmp_path):
        path = tmp_path / "empty.md"
        path.write_text("", encoding="utf-8")

        loader = MarkdownLoader()
        pages = await loader.load(str(path))
        assert len(pages) == 1

    @pytest.mark.asyncio
    async def test_handles_unicode(self, tmp_path):
        path = tmp_path / "unicode.md"
        path.write_text("# 你好世界\n\n日本語テスト。", encoding="utf-8")

        loader = MarkdownLoader()
        pages = await loader.load(str(path))
        assert "你好世界" in pages[0].content

    @pytest.mark.asyncio
    async def test_handles_missing_file(self, tmp_path):
        # _get_file_info is called before the try/except
        loader = MarkdownLoader()
        with pytest.raises(FileNotFoundError):
            await loader.load(str(tmp_path / "nonexistent.md"))

    @pytest.mark.asyncio
    async def test_large_markdown_splits_pages(self, tmp_path):
        path = tmp_path / "big.md"
        path.write_text("x" * 6000, encoding="utf-8")

        loader = MarkdownLoader()
        pages = await loader.load(str(path))
        assert len(pages) >= 2


# ── HTMLLoader ──────────────────────────────────────────────────


class TestHTMLLoader:
    @pytest.mark.asyncio
    async def test_parses_html_and_converts_to_markdown(self, tmp_path):
        path = tmp_path / "page.html"
        path.write_text(
            "<html><body><h1>Hello</h1><p>This is a paragraph.</p></body></html>",
            encoding="utf-8",
        )

        loader = HTMLLoader()
        pages = await loader.load(str(path))

        assert len(pages) >= 1
        assert "Hello" in pages[0].content
        assert "paragraph" in pages[0].content

    @pytest.mark.asyncio
    async def test_strips_script_and_style_tags(self, tmp_path):
        path = tmp_path / "with_scripts.html"
        path.write_text(
            "<html><head><script>alert('xss')</script><style>body{color:red}</style></head>"
            "<body><p>Visible content.</p></body></html>",
            encoding="utf-8",
        )

        loader = HTMLLoader()
        pages = await loader.load(str(path))
        content = " ".join(p.content for p in pages)
        assert "alert" not in content
        assert "xss" not in content
        assert "color" not in content
        assert "Visible content" in content

    @pytest.mark.asyncio
    async def test_handles_malformed_html(self, tmp_path):
        path = tmp_path / "malformed.html"
        path.write_text("<p>Unclosed paragraph<span>text<p>another", encoding="utf-8")

        loader = HTMLLoader()
        pages = await loader.load(str(path))
        assert len(pages) >= 1
        assert "Unclosed paragraph" in pages[0].content

    @pytest.mark.asyncio
    async def test_handles_empty_html_file(self, tmp_path):
        path = tmp_path / "empty.html"
        path.write_text("<html><body></body></html>", encoding="utf-8")

        loader = HTMLLoader()
        pages = await loader.load(str(path))
        assert len(pages) == 1

    @pytest.mark.asyncio
    async def test_handles_missing_file(self, tmp_path):
        # _get_file_info is called before inner try block
        loader = HTMLLoader()
        with pytest.raises(FileNotFoundError):
            await loader.load(str(tmp_path / "nonexistent.html"))

    @pytest.mark.asyncio
    async def test_strips_nav_and_footer_tags(self, tmp_path):
        path = tmp_path / "with_nav.html"
        path.write_text(
            "<html><body>"
            "<nav>Home | About</nav>"
            "<main>Main content here.</main>"
            "<footer>Copyright 2023</footer>"
            "</body></html>",
            encoding="utf-8",
        )

        loader = HTMLLoader()
        pages = await loader.load(str(path))
        content = " ".join(p.content for p in pages)
        assert "Home" not in content
        assert "Copyright" not in content
        assert "Main content" in content

    @pytest.mark.asyncio
    async def test_handles_html_with_special_chars(self, tmp_path):
        path = tmp_path / "special.html"
        path.write_text(
            "<html><body><p>&lt;tag&gt; &amp; &copy; &euro;</p></body></html>",
            encoding="utf-8",
        )

        loader = HTMLLoader()
        pages = await loader.load(str(path))
        assert len(pages) >= 1

    @pytest.mark.asyncio
    async def test_html_without_body_tag(self, tmp_path):
        path = tmp_path / "no_body.html"
        path.write_text("<h1>Title</h1><p>Content without body tag.</p>", encoding="utf-8")

        loader = HTMLLoader()
        pages = await loader.load(str(path))
        assert "Title" in pages[0].content
        assert "Content without body tag" in pages[0].content


# ── _clean_pdf_text edge cases ──────────────────────────────────


class TestCleanPdfTextEdgeCases:
    def test_empty_string_returns_empty(self):
        assert _clean_pdf_text("") == ""

    def test_only_whitespace(self):
        result = _clean_pdf_text("   \t  \n\n\n  ")
        assert result.strip() == ""

    def test_multiple_cid_patterns(self):
        text = "(cid:1)(cid:22)(cid:333) text"
        result = _clean_pdf_text(text)
        assert "(cid:" not in result
        assert "text" in result

    def test_cid_within_text(self):
        text = "Hello(cid:99)World"
        result = _clean_pdf_text(text)
        assert "(cid:99)" not in result
        assert "HelloWorld" in result

    def test_noise_phrase_removed(self):
        # Regex matches 散不轻把 + 1-4 non-whitespace chars (greedy 4)
        # "散不轻把成银继续" → all removed, leaving just the prefix
        text = "正文散不轻把成银继续"
        result = _clean_pdf_text(text)
        assert "散不轻把" not in result
        assert "正文" in result

    def test_noise_phrase_at_start(self):
        # "散不轻把成银正文" (4+4 chars) removed, leaving "开始"
        text = "散不轻把成银正文开始"
        result = _clean_pdf_text(text)
        assert "散不轻把" not in result
        assert "开始" in result

    def test_normalizes_whitespace(self):
        text = "hello    world\t\tfoo\n\n\n\nbar"
        result = _clean_pdf_text(text)
        assert "hello world foo" in result
        assert result.count("\n\n\n") == 0

    def test_non_string_inputs_like_int(self):
        result = _clean_pdf_text("12345")
        assert result == "12345"

    def test_none_like_string_literal(self):
        result = _clean_pdf_text("None")
        assert result == "None"

    def test_only_newlines(self):
        result = _clean_pdf_text("\n\n\n\n\n")
        assert result.strip() == ""

    def test_mixed_cid_and_noise(self):
        text = "散不轻把成银(cid:777) 实际文本"
        result = _clean_pdf_text(text)
        assert "散不轻把成银" not in result
        assert "(cid:777)" not in result
        assert "实际文本" in result
