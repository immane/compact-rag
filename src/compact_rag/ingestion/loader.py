from __future__ import annotations

import hashlib
import os
import re
from abc import ABC, abstractmethod
from pathlib import Path
from typing import ClassVar

from compact_rag.common.exceptions import DocumentLoadError, UnsupportedFormatError
from compact_rag.common.logger import get_logger

logger = get_logger(__name__)


class LoadedPage:
    """A single page of loaded document content."""

    def __init__(
        self,
        page_number: int,
        content: str,
        tables: list[str] | None = None,
        metadata: dict | None = None,
    ) -> None:
        self.page_number = page_number
        self.content = content
        self.tables = tables or []
        self.metadata = metadata or {}


class BaseLoader(ABC):
    """Abstract base for document loaders."""

    @abstractmethod
    async def load(self, file_path: str) -> list[LoadedPage]:
        """Load a document and return its pages."""

    @staticmethod
    def _get_file_info(file_path: str) -> dict:
        path = Path(file_path)
        file_size = path.stat().st_size
        sha256_hash = hashlib.sha256(path.read_bytes()).hexdigest()
        ext = path.suffix.lower()
        ext_map = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".txt": "txt",
            ".md": "md",
            ".html": "html",
            ".htm": "html",
        }
        return {
            "filename": path.name,
            "file_type": ext_map.get(ext, ext.lstrip(".")),
            "file_size": file_size,
            "hash": sha256_hash,
        }


class PDFLoader(BaseLoader):
    """PDF document loader using pypdf."""

    async def load(self, file_path: str) -> list[LoadedPage]:
        file_info = self._get_file_info(file_path)
        pages: list[LoadedPage] = []

        try:
            extracted = await self._extract_with_pdfplumber(file_path)
            if extracted is None:
                extracted = await self._extract_with_pypdf(file_path)

            page_count = len(extracted)
            for i, text in enumerate(extracted):
                pages.append(
                    LoadedPage(
                        page_number=i + 1,
                        content=_clean_pdf_text(text).strip(),
                        tables=[],
                        metadata={"page_count": page_count, **file_info},
                    )
                )
        except Exception as e:
            raise DocumentLoadError(
                f"Failed to load PDF '{os.path.basename(file_path)}': {e}", cause=e
            )

        logger.info("PDF loaded", filename=file_info["filename"], pages=len(pages))
        return pages

    async def _extract_with_pdfplumber(self, file_path: str) -> list[str] | None:
        try:
            import pdfplumber
        except ImportError:
            return None

        texts: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                texts.append(page.extract_text() or "")
        return texts

    async def _extract_with_pypdf(self, file_path: str) -> list[str]:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise DocumentLoadError("pypdf not installed. Install with: pip install pypdf")

        reader = PdfReader(file_path)
        return [(page.extract_text() or "") for page in reader.pages]


_PDF_NOISE_PATTERNS = [
    re.compile(r"散不轻把[^\s，。,｡；;:：]{1,4}"),
    re.compile(r"\(cid:\d+\)"),
]


def _clean_pdf_text(text: str) -> str:
    cleaned = text
    for pattern in _PDF_NOISE_PATTERNS:
        cleaned = pattern.sub("", cleaned)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned


class DOCXLoader(BaseLoader):
    """DOCX document loader using python-docx."""

    async def load(self, file_path: str) -> list[LoadedPage]:
        try:
            from docx import Document
        except ImportError:
            raise DocumentLoadError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        file_info = self._get_file_info(file_path)
        pages: list[LoadedPage] = []

        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables: list[str] = []

            for table in doc.tables:
                md = _docx_table_to_markdown(table)
                if md:
                    tables.append(md)

            full_text = "\n\n".join(paragraphs)
            page_size = 3000
            page_parts = [
                full_text[i : i + page_size]
                for i in range(0, len(full_text), page_size)
            ]

            for i, part in enumerate(page_parts):
                pages.append(
                    LoadedPage(
                        page_number=i + 1,
                        content=part.strip(),
                        tables=tables if i == 0 else [],
                        metadata={
                            "page_count": len(page_parts),
                            "table_count": len(tables),
                            **file_info,
                        },
                    )
                )

            if not pages:
                pages.append(
                    LoadedPage(
                        page_number=1,
                        content=full_text.strip(),
                        tables=tables,
                        metadata={
                            "page_count": 1,
                            "table_count": len(tables),
                            **file_info,
                        },
                    )
                )
        except Exception as e:
            raise DocumentLoadError(
                f"Failed to load DOCX '{os.path.basename(file_path)}': {e}", cause=e
            )

        logger.info("DOCX loaded", filename=file_info["filename"], pages=len(pages))
        return pages


def _docx_table_to_markdown(table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    col_count = max(len(r) for r in rows)
    for r in rows:
        while len(r) < col_count:
            r.append("")
    lines: list[str] = []
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * col_count) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


class TextLoader(BaseLoader):
    """Plain text document loader."""

    async def load(self, file_path: str) -> list[LoadedPage]:
        file_info = self._get_file_info(file_path)
        try:
            content = Path(file_path).read_text(encoding="utf-8")
        except UnicodeDecodeError:
            content = Path(file_path).read_text(encoding="latin-1")
        except Exception as e:
            raise DocumentLoadError(
                f"Failed to load text file '{os.path.basename(file_path)}': {e}", cause=e
            )

        page_size = 3000
        page_parts = [
            content[i : i + page_size]
            for i in range(0, len(content), page_size)
        ]

        pages: list[LoadedPage] = []
        for i, part in enumerate(page_parts):
            pages.append(
                LoadedPage(
                    page_number=i + 1,
                    content=part.strip(),
                    metadata={"page_count": len(page_parts) or 1, **file_info},
                )
            )
        if not pages:
            pages.append(
                LoadedPage(
                    page_number=1,
                    content=content.strip(),
                    metadata={"page_count": 1, **file_info},
                )
            )
        logger.info("Text loaded", filename=file_info["filename"], pages=len(pages))
        return pages


class MarkdownLoader(BaseLoader):
    """Markdown document loader."""

    async def load(self, file_path: str) -> list[LoadedPage]:
        file_info = self._get_file_info(file_path)
        try:
            content = Path(file_path).read_text(encoding="utf-8")
        except Exception as e:
            raise DocumentLoadError(
                f"Failed to load markdown file '{os.path.basename(file_path)}': {e}", cause=e
            )

        page_size = 3000
        page_parts = [
            content[i : i + page_size]
            for i in range(0, len(content), page_size)
        ]

        pages: list[LoadedPage] = []
        for i, part in enumerate(page_parts):
            pages.append(
                LoadedPage(
                    page_number=i + 1,
                    content=part.strip(),
                    metadata={"page_count": len(page_parts) or 1, **file_info},
                )
            )
        if not pages:
            pages.append(
                LoadedPage(
                    page_number=1,
                    content=content.strip(),
                    metadata={"page_count": 1, **file_info},
                )
            )
        logger.info("Markdown loaded", filename=file_info["filename"], pages=len(pages))
        return pages


class HTMLLoader(BaseLoader):
    """HTML document loader using BeautifulSoup + markdownify."""

    async def load(self, file_path: str) -> list[LoadedPage]:
        try:
            from bs4 import BeautifulSoup
            from markdownify import markdownify
        except ImportError:
            raise DocumentLoadError(
                "bs4/markdownify not installed. Install with: pip install beautifulsoup4 markdownify"
            )

        file_info = self._get_file_info(file_path)
        try:
            html_content = Path(file_path).read_text(encoding="utf-8")
        except Exception as e:
            raise DocumentLoadError(
                f"Failed to load HTML file '{os.path.basename(file_path)}': {e}", cause=e
            )

        try:
            soup = BeautifulSoup(html_content, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            body = soup.find("body") or soup
            md_content = markdownify(str(body), heading_style="ATX", strip=["a", "img"])

            page_size = 3000
            page_parts = [
                md_content[i : i + page_size]
                for i in range(0, len(md_content), page_size)
            ]

            pages: list[LoadedPage] = []
            for i, part in enumerate(page_parts):
                pages.append(
                    LoadedPage(
                        page_number=i + 1,
                        content=part.strip(),
                        metadata={"page_count": len(page_parts) or 1, **file_info},
                    )
                )
            if not pages:
                pages.append(
                    LoadedPage(
                        page_number=1,
                        content=md_content.strip(),
                        metadata={"page_count": 1, **file_info},
                    )
                )
        except Exception as e:
            raise DocumentLoadError(
                f"Failed to parse HTML '{os.path.basename(file_path)}': {e}", cause=e
            )

        logger.info("HTML loaded", filename=file_info["filename"], pages=len(pages))
        return pages


class LoaderFactory:
    """Factory that returns the appropriate loader based on file extension."""

    _loaders: ClassVar[dict[str, type[BaseLoader]]] = {
        ".pdf": PDFLoader,
        ".docx": DOCXLoader,
        ".txt": TextLoader,
        ".md": MarkdownLoader,
        ".html": HTMLLoader,
        ".htm": HTMLLoader,
    }

    @classmethod
    def get_loader(cls, file_path: str) -> BaseLoader:
        ext = Path(file_path).suffix.lower()
        loader_class = cls._loaders.get(ext)
        if loader_class is None:
            raise UnsupportedFormatError(
                f"Unsupported file format: '{ext}'. Supported: {list(cls._loaders.keys())}"
            )
        return loader_class()
